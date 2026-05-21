import sys
import os
import io
import json
import time

sys.path.append(os.path.dirname(os.path.dirname(__file__)))

from fastapi import FastAPI, UploadFile, File, HTTPException, Body, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, Response
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from jose import JWTError, jwt
import hashlib
from pydantic import BaseModel
from openai import AsyncOpenAI
import asyncio
from datetime import datetime, timedelta
from urllib.parse import quote

from backend.core import (
    run_expert_analysis,
    run_summary_analysis,
    run_expert_analysis_stream,
    run_summary_analysis_stream,
    PROMPT_CONFIGS,
)
from backend.db import (
    init_db,
    get_db,
    execute_query,
    get_db_type,
)
from backend.config import load_config, save_config, DEFAULT_SECRET_KEY
from backend.pdf_generator import generate_pdf_report
from backend.text_extractor import extract_text
from backend.keyword_extractor import extract_keywords
from fastapi.responses import FileResponse
import docx as docx_module

from backend.logger import log_api, log_analysis, log_llm, log_error

app = FastAPI()


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def verify_password(plain_password: str, hashed_password: str) -> bool:
    hashed_input = hash_password(plain_password)
    return hashed_input == hashed_password


SECRET_KEY = load_config().get("secretKey", DEFAULT_SECRET_KEY)
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/token")


class User(BaseModel):
    username: str
    email: str = None
    role: str = "user"


class UserCreate(BaseModel):
    username: str
    password: str
    email: str = None


class Token(BaseModel):
    access_token: str
    token_type: str


class TokenData(BaseModel):
    username: str = None


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

analysis_tasks = {}

init_db()


def _log(tag, msg):
    ts = time.strftime("%H:%M:%S")
    print(f"  \033[90m[{ts}]\033[0m {tag} {msg}")


@app.middleware("http")
async def logging_middleware(request: Request, call_next):
    skip_paths = ("/api/analysis/status", "/favicon.ico")
    if any(request.url.path.startswith(p) for p in skip_paths):
        return await call_next(request)
    start = time.time()
    response = await call_next(request)
    duration = (time.time() - start) * 1000
    log_api(request.method, request.url.path, response.status_code, duration)
    return response


@app.on_event("startup")
async def startup():
    _log("\033[36m[API]\033[0m", "\033[32m✓ FastAPI 路由注册完成\033[0m")
    _log("\033[36m[DB]\033[0m", "\033[32m✓ 数据库初始化完成\033[0m")
    _log("\033[36m[API]\033[0m", f"  已注册端点: \033[33m/api/analyze/stream\033[0m | \033[33m/api/extract-keywords\033[0m | \033[33m/api/download/pdf\033[0m | \033[33m/api/download/docx\033[0m")


def row_to_dict(cursor, row):
    if row and isinstance(row, tuple) and cursor.description:
        columns = [desc[0] for desc in cursor.description]
        return dict(zip(columns, row))
    return row


def rows_to_dicts(cursor, rows):
    if not rows or not cursor.description:
        return rows
    columns = [desc[0] for desc in cursor.description]
    return [dict(zip(columns, row)) for row in rows]


@app.get("/favicon.ico")
async def favicon():
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.dirname(__file__)))
    path = os.path.join(base, "favicon.ico")
    if os.path.exists(path):
        return FileResponse(path)
    raise HTTPException(status_code=404)


def get_user(username: str):
    with get_db() as conn:
        cursor = execute_query(
            conn, "SELECT * FROM users WHERE username = ?", (username,)
        )
        return row_to_dict(cursor, cursor.fetchone())


def authenticate_user(username: str, password: str):
    user = get_user(username)
    if not user:
        return False
    if not verify_password(password, user["password"]):
        return False
    return user


def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt


async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception
    user = get_user(username=token_data.username)
    if user is None:
        raise credentials_exception

    # 检查token是否与数据库中存储的token一致
    if user.get("current_token") != token:
        raise credentials_exception

    # 更新最后活动时间
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with get_db() as conn:
            execute_query(
                conn,
                "UPDATE users SET last_activity = ? WHERE username = ?",
                (now, username),
            )
            conn.commit()
    except Exception:
        pass

    return user


async def get_current_active_user(current_user: dict = Depends(get_current_user)):
    return current_user


def requires_auth():
    config = load_config()
    return config.get("requireAuth", True)


async def optional_auth(request: Request):
    if requires_auth():
        token = request.headers.get("Authorization")
        if not token:
            raise HTTPException(
                status_code=401,
                detail="Could not validate credentials",
                headers={"WWW-Authenticate": "Bearer"},
            )
        token = token.replace("Bearer ", "")
        return await get_current_user(token)
    return None


def cleanup_completed_tasks(max_keep=10):
    completed_tasks = [
        tid
        for tid, task in analysis_tasks.items()
        if task.get("status") in ("completed", "error")
    ]
    if len(completed_tasks) > max_keep:
        for tid in completed_tasks[:-max_keep]:
            analysis_tasks.pop(tid, None)


@app.post("/api/analyze")
async def analyze(
    content: str = Body(...),
    api_key: str = Body(...),
    base_url: str = Body(...),
    configs: dict = Body(None),
    current_user: dict = Depends(optional_auth),
):
    if not content:
        raise HTTPException(status_code=400, detail="Content is required")

    if not api_key or not base_url:
        raise HTTPException(status_code=400, detail="API Key and Base URL are required")

    all_configs = configs if configs else PROMPT_CONFIGS

    cleanup_completed_tasks()

    task_id = f"task_{datetime.now().timestamp()}"
    analysis_tasks[task_id] = {
        "status": "started",
        "progress": 10,
        "current_step": "正在初始化分析任务...",
        "reports": [],
        "final_report": None,
    }

    async def perform_analysis():
        try:
            client = AsyncOpenAI(api_key=api_key, base_url=base_url)

            analysis_tasks[task_id]["status"] = "running_experts"
            analysis_tasks[task_id]["progress"] = 15
            analysis_tasks[task_id]["current_step"] = (
                "⚖️ 刑法专家 📋 合规审查专家 🔎 证据分析专家 并行分析中..."
            )

            expert_configs = {
                "expert1": all_configs.get("expert1", all_configs["expert1"]),
                "expert2": all_configs.get("expert2", all_configs["expert2"]),
                "expert3": all_configs.get("expert3", all_configs["expert3"]),
            }

            async def update_progress_during_wait():
                for i in range(3):
                    await asyncio.sleep(8)
                    analysis_tasks[task_id]["progress"] = 25 + i * 8
                    analysis_tasks[task_id]["current_step"] = (
                        f"⚖️ 刑法专家 📋 合规审查专家 🔎 证据分析专家 工作中... ({i + 1}/3)"
                    )

            progress_task = asyncio.create_task(update_progress_during_wait())
            reports = await run_expert_analysis(client, content, expert_configs)
            analysis_tasks[task_id]["reports"] = list(reports)
            progress_task.cancel()

            analysis_tasks[task_id]["status"] = "experts_done"
            analysis_tasks[task_id]["progress"] = 50
            analysis_tasks[task_id]["current_step"] = (
                "三位专家分析完成！正在生成汇总报告..."
            )

            analysis_tasks[task_id]["status"] = "running_summary"
            analysis_tasks[task_id]["progress"] = 60
            analysis_tasks[task_id]["current_step"] = "📖 汇总专家 整合意见中..."

            final_report = await run_summary_analysis(
                client, analysis_tasks[task_id]["reports"], all_configs
            )
            analysis_tasks[task_id]["final_report"] = final_report
            analysis_tasks[task_id]["status"] = "completed"
            analysis_tasks[task_id]["progress"] = 100
            analysis_tasks[task_id]["current_step"] = "✅ 分析完成"
        except Exception as e:
            analysis_tasks[task_id]["status"] = "error"
            analysis_tasks[task_id]["error"] = str(e)
            analysis_tasks[task_id]["current_step"] = f"错误: {str(e)}"

    asyncio.create_task(perform_analysis())

    return {"task_id": task_id}


async def event_generator(
    task_id: str, api_key: str, base_url: str, content: str, all_configs: dict
):
    client = AsyncOpenAI(api_key=api_key, base_url=base_url)
    expert_names = [v["name"] for k, v in sorted(all_configs.items()) if k.startswith("expert")]
    expert_configs = {k: v for k, v in sorted(all_configs.items()) if k.startswith("expert")}

    log_analysis(f"━━━ 开始并行流式分析 ━━━  专家: {len(expert_configs)}位  模型: {list(expert_configs.values())[0]['model']}")
    log_analysis(f"  任务ID: {task_id}  输入: {len(content)} 字")
    yield f'data: {{"type": "start", "message": "开始专家分析..."}}\n\n'

    expert_texts = [""] * len(expert_configs)

    async for expert_idx, chunk in run_expert_analysis_stream(client, content, expert_configs):
        if expert_idx == -1 and chunk == "[ALL_DONE]":
            log_analysis(f"  ✅ 全部专家分析完成  共 {len(expert_texts)} 份报告")
            yield f"data: {json.dumps({'type': 'experts_done', 'reports': expert_texts})}\n\n"
            continue
        expert_name = expert_names[expert_idx] if expert_idx < len(expert_names) else f"专家{expert_idx+1}"
        if chunk.startswith("[ERROR]"):
            log_error(f"  ❌ {expert_name} 分析失败: {chunk}")
            yield f"data: {json.dumps({'type': 'error', 'expert': expert_idx, 'message': chunk})}\n\n"
            return
        if chunk.startswith("[FINAL]"):
            expert_texts[expert_idx] = chunk.replace("[FINAL]", "").strip()
            log_analysis(f"  ✅ {expert_name} 报告生成完成  ({len(expert_texts[expert_idx])} 字)")
            yield f"data: {json.dumps({'type': 'expert_done', 'expert': expert_idx, 'name': expert_name, 'text': expert_texts[expert_idx]})}\n\n"
        else:
            expert_texts[expert_idx] += chunk
            yield f"data: {json.dumps({'type': 'expert_stream', 'expert': expert_idx, 'name': expert_name, 'chunk': chunk})}\n\n"

    log_analysis("━━━ 开始汇总分析 ━━━")
    yield f'data: {{"type": "summary_start", "message": "开始汇总分析..."}}\n\n'

    summary_text = ""
    async for chunk in run_summary_analysis_stream(client, expert_texts, all_configs):
        if chunk.startswith("[ERROR]"):
            log_error(f"  ❌ 汇总分析失败: {chunk}")
            yield f"data: {json.dumps({'type': 'error', 'summary': True, 'message': chunk})}\n\n"
            return
        if chunk.startswith("[FINAL]"):
            summary_text = chunk.replace("[FINAL]", "").strip()
            log_analysis(f"  ✅ 汇总报告生成完成  ({len(summary_text)} 字)")
            yield f"data: {json.dumps({'type': 'summary_done', 'report': summary_text})}\n\n"
        else:
            summary_text += chunk
            yield f"data: {json.dumps({'type': 'summary_stream', 'chunk': chunk})}\n\n"

    log_analysis(f"━━━ 分析流程结束 ━━━  总报告: {len(summary_text)} 字")
    yield f"data: {json.dumps({'type': 'complete', 'report': summary_text, 'reports': expert_texts})}\n\n"


@app.post("/api/analyze/stream")
async def analyze_stream(
    content: str = Body(...),
    api_key: str = Body(...),
    base_url: str = Body(...),
    configs: dict = Body(None),
    current_user: dict = Depends(optional_auth),
):
    if not content:
        raise HTTPException(status_code=400, detail="Content is required")
    if not api_key or not base_url:
        raise HTTPException(status_code=400, detail="API Key and Base URL are required")

    all_configs = configs if configs else PROMPT_CONFIGS

    import uuid
    task_id = f"task_{uuid.uuid4().hex[:8]}"
    log_analysis(f"收到分析请求  内容: {len(content)} 字  API: {base_url}  专家: {len([k for k in all_configs if k.startswith('expert')])}位")

    return StreamingResponse(
        event_generator(task_id, api_key, base_url, content, all_configs),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/api/analysis/status/{task_id}")
def get_analysis_status(task_id: str, current_user: dict = Depends(optional_auth)):
    if task_id not in analysis_tasks:
        raise HTTPException(status_code=404, detail="Task not found")

    return analysis_tasks[task_id]


@app.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...), current_user: dict = Depends(optional_auth)
):
    try:
        content = await file.read()
        log_api("FILE", f"文件上传: {file.filename}  类型: {file.content_type}  大小: {len(content)/1024:.1f}KB")
        text = extract_text(content, file.content_type)
        log_api("FILE", f"文本提取完成: {len(text)} 字")
        return {"content": text}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")


@app.post("/api/download/pdf")
def download_pdf(
    final_report: str = Body(...),
    reports: list = Body(...),
    configs: dict = Body(...),
    current_user: dict = Depends(optional_auth),
):
    try:
        log_api("EXPORT", "正在生成 PDF 报告...")
        experts_data = {}
        for i, (key, cfg) in enumerate(configs.items()):
            if i < len(reports):
                experts_data[cfg["name"]] = reports[i]

        pdf_buffer = generate_pdf_report(
            title="法律风险多模型研判报告",
            content=final_report,
            experts_reports=experts_data,
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"法律风险研判报告_{timestamp}.pdf"
        encoded_file_name = quote(file_name)
        log_api("EXPORT", f"PDF 报告生成完成: {file_name}")

        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_file_name}"
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")


@app.post("/api/download/txt")
def download_txt(
    report: str = Body(...),
    expert_name: str = Body(...),
    current_user: dict = Depends(optional_auth),
):
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{expert_name}初步报告_{timestamp}.txt"
        encoded_filename = quote(filename)
        return Response(
            content=report.encode("utf-8"),
            media_type="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating TXT: {str(e)}")


@app.post("/api/download/docx")
def download_docx(
    final_report: str = Body(...),
    reports: list = Body(...),
    configs: dict = Body(...),
    current_user: dict = Depends(optional_auth),
):
    try:
        doc = docx_module.Document()
        doc.add_heading("法律风险多模型研判报告", level=0)
        doc.add_paragraph(final_report)
        for i, (key, cfg) in enumerate(configs.items()):
            if i < len(reports):
                doc.add_heading(cfg["name"], level=1)
                doc.add_paragraph(reports[i])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"法律风险研判报告_{timestamp}.docx"
        encoded_file_name = quote(file_name)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_file_name}"
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


@app.get("/api/config")
def get_config(current_user: dict = Depends(optional_auth)):
    try:
        return load_config()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading config: {str(e)}")


@app.post("/api/config")
def save_config_endpoint(
    config: dict = Body(...), current_user: dict = Depends(optional_auth)
):
    try:
        save_config(config)
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving config: {str(e)}")


@app.get("/api/history")
def get_history(limit: int = 20, current_user: dict = Depends(optional_auth)):
    try:
        with get_db() as conn:
            # 如果有用户，只显示该用户的报告
            if current_user:
                cursor = execute_query(
                    conn,
                    "SELECT id, created_at, title, input_content FROM reports WHERE user_id = ? ORDER BY created_at DESC LIMIT ?",
                    (
                        current_user["id"],
                        limit,
                    ),
                )
            else:
                cursor = execute_query(
                    conn,
                    "SELECT id, created_at, title, input_content FROM reports ORDER BY created_at DESC LIMIT ?",
                    (limit,),
                )
            rows = cursor.fetchall()
            rows = rows_to_dicts(cursor, rows)
        return [
            {
                "id": r["id"],
                "created_at": r["created_at"],
                "title": r["title"]
                or f"法律风险评估报告-{r['created_at'][:10].replace('-', '')}",
                "input_content": r["input_content"][:100] if r["input_content"] else "",
            }
            for r in rows
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading history: {str(e)}")


@app.get("/api/history/{report_id}")
def get_history_detail(report_id: int, current_user: dict = Depends(optional_auth)):
    try:
        with get_db() as conn:
            cursor = execute_query(
                conn,
                """
                SELECT r.*, u.username 
                FROM reports r 
                LEFT JOIN users u ON r.user_id = u.id 
                WHERE r.id = ?
            """,
                (report_id,),
            )
            row = row_to_dict(cursor, cursor.fetchone())
        if not row:
            raise HTTPException(status_code=404, detail="Report not found")
        return row
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading report: {str(e)}")


@app.post("/api/history/save")
def save_history(data: dict = Body(...), current_user: dict = Depends(optional_auth)):
    try:
        created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        title = data.get(
            "title", f"法律风险评估报告-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        )
        user_id = current_user["id"] if current_user else None
        with get_db() as conn:
            execute_query(
                conn,
                """
                INSERT INTO reports (user_id, created_at, title, input_content, final_report, expert1_report, expert2_report, expert3_report)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
                (
                    user_id,
                    created_at,
                    title,
                    data.get("input_content"),
                    data.get("final_report"),
                    data.get("expert1_report"),
                    data.get("expert2_report"),
                    data.get("expert3_report"),
                ),
            )
            conn.commit()
            # 获取最后插入的 ID，根据数据库类型使用不同语法
            if get_db_type() == "mysql":
                report_id = execute_query(conn, "SELECT LAST_INSERT_ID()").fetchone()[0]
            else:
                report_id = execute_query(
                    conn, "SELECT last_insert_rowid()"
                ).fetchone()[0]
        return {"success": True, "id": report_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving history: {str(e)}")


@app.delete("/api/history/{report_id}")
def delete_history(report_id: int, current_user: dict = Depends(optional_auth)):
    try:
        with get_db() as conn:
            execute_query(conn, "DELETE FROM reports WHERE id = ?", (report_id,))
            conn.commit()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting history: {str(e)}")


@app.put("/api/history/{report_id}/title")
def update_history_title(
    report_id: int, data: dict = Body(...), current_user: dict = Depends(optional_auth)
):
    try:
        with get_db() as conn:
            execute_query(
                conn,
                "UPDATE reports SET title = ? WHERE id = ?",
                (data.get("title"), report_id),
            )
            conn.commit()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating title: {str(e)}")


@app.post("/api/extract-keywords")
def get_keywords(data: dict = Body(...), current_user: dict = Depends(optional_auth)):
    try:
        text = data.get("text", "")
        top_n = data.get("top_n", 50)
        keywords = extract_keywords(text, top_n)
        log_api("NLP", f"关键词提取: 输入 {len(text)} 字  提取 {len(keywords)} 个关键词  top_n={top_n}")
        return {"keywords": keywords}
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error extracting keywords: {str(e)}"
        )


@app.post("/api/register", response_model=User)
def register(user: UserCreate):
    try:
        existing_user = get_user(user.username)
        if existing_user:
            log_api("AUTH", f"注册失败: 用户名 {user.username} 已存在", 400)
            raise HTTPException(status_code=400, detail="Username already registered")

        hashed_password = hash_password(user.password)
        created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        with get_db() as conn:
            execute_query(
                conn,
                "INSERT INTO users (username, password, email, created_at) VALUES (?, ?, ?, ?)",
                (user.username, hashed_password, user.email, created_at),
            )
            conn.commit()

        log_api("AUTH", f"新用户注册: {user.username}", 200)
        return {"username": user.username, "email": user.email, "role": "user"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error registering user: {str(e)}")


@app.post("/api/token", response_model=Token)
def login(form_data: OAuth2PasswordRequestForm = Depends()):
    try:
        user = authenticate_user(form_data.username, form_data.password)
        if not user:
            log_api("AUTH", f"登录失败: {form_data.username} 用户名或密码错误", 401)
            raise HTTPException(
                status_code=401,
                detail="Incorrect username or password",
                headers={"WWW-Authenticate": "Bearer"},
            )

        # 检查用户是否已经有活跃会话
        if user.get("current_token"):
            # 检查会话是否过期（3分钟无活动）
            last_activity = user.get("last_activity")
            if last_activity:
                try:
                    last_activity_time = datetime.strptime(
                        last_activity, "%Y-%m-%d %H:%M:%S"
                    )
                    if (datetime.now() - last_activity_time).total_seconds() < 3 * 60:
                        raise HTTPException(
                            status_code=403,
                            detail="User is already logged in from another device",
                        )
                except HTTPException:
                    raise
                except Exception:
                    pass

        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": user["username"]}, expires_delta=access_token_expires
        )

        # 更新最后登录时间、最后活动时间和当前token
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with get_db() as conn:
            execute_query(
                conn,
                "UPDATE users SET last_login = ?, last_activity = ?, current_token = ? WHERE username = ?",
                (now, now, access_token, user["username"]),
            )
            conn.commit()

        log_api("AUTH", f"用户登录: {user['username']}  角色: {user.get('role', 'user')}", 200)
        return {"access_token": access_token, "token_type": "bearer"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error logging in: {str(e)}")


@app.get("/api/users/me", response_model=User)
async def read_users_me(current_user: dict = Depends(get_current_active_user)):
    return {
        "username": current_user["username"],
        "email": current_user.get("email"),
        "role": current_user.get("role", "user"),
    }


@app.post("/api/logout")
async def logout(current_user: dict = Depends(get_current_active_user)):
    """登出API，清除数据库中的当前token"""
    try:
        with get_db() as conn:
            execute_query(
                conn,
                "UPDATE users SET current_token = NULL WHERE username = ?",
                (current_user["username"],),
            )
            conn.commit()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error logging out: {str(e)}")


@app.get("/api/users")
async def get_users(current_user: dict = Depends(get_current_active_user)):
    try:
        if current_user.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Not enough permissions")

        with get_db() as conn:
            cursor = execute_query(
                conn,
                "SELECT id, username, email, created_at, last_login, role FROM users",
            )
            users = cursor.fetchall()
            users = rows_to_dicts(cursor, users)

        return users
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting users: {str(e)}")


@app.delete("/api/users/{user_id}")
async def delete_user(
    user_id: int, current_user: dict = Depends(get_current_active_user)
):
    try:
        if current_user.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Not enough permissions")

        with get_db() as conn:
            execute_query(conn, "DELETE FROM users WHERE id = ?", (user_id,))
            conn.commit()

        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting user: {str(e)}")


class UserConfigUpdate(BaseModel):
    api_key: str = None
    base_url: str = None
    model: str = None
    prompt_configs: dict = None


@app.get("/api/users/{username}/config")
async def get_user_config(
    username: str, current_user: dict = Depends(get_current_active_user)
):
    try:
        # 只能获取自己的配置，或者管理员可以获取任意用户的配置
        if current_user.get("role") != "admin" and current_user["username"] != username:
            raise HTTPException(status_code=403, detail="Not enough permissions")

        # 获取用户
        target_user = get_user(username)
        if not target_user:
            raise HTTPException(status_code=404, detail="User not found")

        with get_db() as conn:
            cursor = execute_query(
                conn,
                "SELECT * FROM user_configs WHERE user_id = ?",
                (target_user["id"],),
            )
            config = cursor.fetchone()
            config = row_to_dict(cursor, config)

            if not config:
                return {
                    "api_key": None,
                    "base_url": None,
                    "model": None,
                    "prompt_configs": None,
                }

            return {
                "api_key": config["api_key"],
                "base_url": config["base_url"],
                "model": config["model"],
                "prompt_configs": json.loads(config["prompt_configs"])
                if config["prompt_configs"]
                else None,
            }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error getting user config: {str(e)}"
        )


@app.post("/api/users/{username}/config")
async def save_user_config(
    username: str,
    config: UserConfigUpdate,
    current_user: dict = Depends(get_current_active_user),
):
    try:
        # 只能修改自己的配置，或者管理员可以修改任意用户的配置
        if current_user.get("role") != "admin" and current_user["username"] != username:
            raise HTTPException(status_code=403, detail="Not enough permissions")

        # 获取用户
        target_user = get_user(username)
        if not target_user:
            raise HTTPException(status_code=404, detail="User not found")

        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        with get_db() as conn:
            # 检查是否已存在配置
            cursor = execute_query(
                conn,
                "SELECT id FROM user_configs WHERE user_id = ?",
                (target_user["id"],),
            )
            existing = cursor.fetchone()

            prompt_configs_json = (
                json.dumps(config.prompt_configs) if config.prompt_configs else None
            )

            if existing:
                execute_query(
                    conn,
                    """
                    UPDATE user_configs 
                    SET api_key = ?, base_url = ?, model = ?, prompt_configs = ?, updated_at = ?
                    WHERE user_id = ?
                    """,
                    (
                        config.api_key,
                        config.base_url,
                        config.model,
                        prompt_configs_json,
                        now,
                        target_user["id"],
                    ),
                )
            else:
                execute_query(
                    conn,
                    """
                    INSERT INTO user_configs (user_id, api_key, base_url, model, prompt_configs, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        target_user["id"],
                        config.api_key,
                        config.base_url,
                        config.model,
                        prompt_configs_json,
                        now,
                        now,
                    ),
                )
            conn.commit()

        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error saving user config: {str(e)}"
        )


class UserUpdate(BaseModel):
    email: str = None
    role: str = None
    password: str = None


@app.put("/api/users/{user_id}")
async def update_user(
    user_id: int,
    user_data: UserUpdate,
    current_user: dict = Depends(get_current_active_user),
):
    try:
        if current_user.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Not enough permissions")

        with get_db() as conn:
            # 检查用户是否存在
            cursor = execute_query(
                conn, "SELECT id FROM users WHERE id = ?", (user_id,)
            )
            if not cursor.fetchone():
                raise HTTPException(status_code=404, detail="User not found")

            # 构建更新语句
            updates = []
            params = []

            if user_data.email is not None:
                updates.append("email = ?")
                params.append(user_data.email)
            if user_data.role is not None:
                updates.append("role = ?")
                params.append(user_data.role)
            if user_data.password is not None:
                updates.append("password = ?")
                params.append(hash_password(user_data.password))

            if updates:
                params.append(user_id)
                execute_query(
                    conn,
                    f"UPDATE users SET {', '.join(updates)} WHERE id = ?",
                    tuple(params),
                )
                conn.commit()

        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating user: {str(e)}")


# 管理员查看所有报告
@app.get("/api/admin/reports")
async def get_all_reports(current_user: dict = Depends(get_current_active_user)):
    try:
        if current_user.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Not enough permissions")

        with get_db() as conn:
            cursor = execute_query(
                conn,
                """
                SELECT r.*, u.username 
                FROM reports r 
                LEFT JOIN users u ON r.user_id = u.id 
                ORDER BY r.created_at DESC
                """,
            )
            rows = rows_to_dicts(cursor, cursor.fetchall())

        return rows
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error getting all reports: {str(e)}"
        )
